import io
import logging
from dataclasses import dataclass
from typing import List
import magic
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


@dataclass
class PageText:
    page: int
    text: str
    start_char: int
    end_char: int


@dataclass
class ParsingResult:
    text: str
    mime_type: str
    parser: str
    page_count: int | None = None
    ocr_used: bool = False
    pages: List[PageText] | None = None


def extract_text(filename: str, data: bytes) -> str:
    return extract_document(filename, data).text


def extract_document(
    filename: str,
    data: bytes,
    enable_ocr: bool = False,
    min_text_chars_for_ocr: int = 100,
) -> ParsingResult:
    mime = magic.from_buffer(data, mime=True)
    lower_name = filename.lower()

    if mime == "application/pdf" or lower_name.endswith(".pdf"):
        text, page_count, pages = _extract_pdf(data)
        if enable_ocr and len(text.strip()) < min_text_chars_for_ocr:
            ocr_text = _extract_pdf_ocr(data)
            if ocr_text.strip():
                return ParsingResult(
                    text=ocr_text,
                    mime_type=mime,
                    parser="pdf_ocr",
                    page_count=page_count,
                    ocr_used=True,
                )
        return ParsingResult(text=text, mime_type=mime, parser="pdfplumber", page_count=page_count, pages=pages)
    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or lower_name.endswith((".docx", ".doc")):
        text = _extract_docx(data)
        return ParsingResult(text=text, mime_type=mime, parser="python-docx")
    if mime.startswith("text/") or lower_name.endswith((".txt", ".md", ".csv", ".json")):
        return ParsingResult(text=data.decode("utf-8", errors="ignore"), mime_type=mime, parser="text")
    if mime == "text/html" or lower_name.endswith((".html", ".htm")):
        return ParsingResult(text=_extract_html(data), mime_type=mime, parser="beautifulsoup")

    # Fallback: prova come testo
    return ParsingResult(text=data.decode("utf-8", errors="ignore"), mime_type=mime, parser="text_fallback")


def _render_markdown_table(rows: List[List[str | None]]) -> str:
    cleaned: List[List[str]] = []
    for row in rows:
        cleaned_row = [(cell or "").strip().replace("|", "\\|").replace("\n", " ") for cell in row]
        if any(cleaned_row):
            cleaned.append(cleaned_row)
    if not cleaned:
        return ""
    width = max(len(row) for row in cleaned)
    cleaned = [row + [""] * (width - len(row)) for row in cleaned]
    header, *body = cleaned
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


_TABLE_SETTINGS = {"vertical_strategy": "lines_strict", "horizontal_strategy": "lines_strict"}


def _extract_pdf_page(page) -> str:
    # Le tabelle vengono estratte come Markdown ed escluse dal testo libero per evitare
    # di duplicarne il contenuto in forma illeggibile (celle appiattite su una riga).
    # Si richiedono linee di confine reali (lines_strict): le strategie di default di
    # pdfplumber ("lines"/"text") inferiscono colonne dalla posizione del testo e producono
    # falsi positivi sistematici su documenti di testo denso senza una vera griglia tabellare
    # (verificato su un dossier parlamentare reale: rilevava "tabelle" da paragrafi di prosa).
    tables = page.find_tables(table_settings=_TABLE_SETTINGS)
    text_page = page
    for table in tables:
        text_page = text_page.outside_bbox(table.bbox)
    parts = []
    prose = (text_page.extract_text() or "").strip()
    if prose:
        parts.append(prose)
    for index, table in enumerate(tables, start=1):
        rendered = _render_markdown_table(table.extract())
        if rendered:
            parts.append(f"Tabella {index}:\n\n{rendered}")
    if page.images:
        # Niente modello vision: segnaliamo solo la presenza di contenuto non testuale.
        plural = "i" if len(page.images) > 1 else ""
        parts.append(f"[Pagina con {len(page.images)} immagine{plural}/figura{plural} non elaborata{plural}]")
    return "\n\n".join(parts)


def _extract_pdf(data: bytes) -> tuple[str, int, List[PageText]]:
    parts: List[str] = []
    pages: List[PageText] = []
    cursor = 0
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for page_index, page in enumerate(pdf.pages, start=1):
            text = _extract_pdf_page(page)
            if text:
                start_char = cursor
                parts.append(text)
                cursor += len(text)
                pages.append(PageText(page=page_index, text=text, start_char=start_char, end_char=cursor))
                cursor += 2
    return "\n\n".join(parts), page_count, pages


def _extract_pdf_ocr(data: bytes) -> str:
    try:
        import pypdfium2 as pdfium  # type: ignore
        import pytesseract  # type: ignore
    except Exception as exc:
        logger.warning("OCR requested but optional OCR dependencies are unavailable: %s", exc)
        return ""

    parts: List[str] = []
    pdf = pdfium.PdfDocument(data)
    try:
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            bitmap = page.render(scale=2).to_pil()
            text = pytesseract.image_to_string(bitmap)
            if text.strip():
                parts.append(text)
    finally:
        pdf.close()
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    return "\n\n".join(parts)


def _extract_html(data: bytes) -> str:
    soup = BeautifulSoup(data, "html.parser")
    for script in soup(["script", "style"]):
        script.decompose()
    return soup.get_text(separator="\n")
